"""
Sleep Lab Tech Note App
------------------------
A simple desktop GUI (Tkinter) for entering sleep-study tech note data
and exporting it to a formatted Word (.docx) document.

Dependencies (install with pip):
    pip install tkcalendar python-docx

Run with:
    python tech_note_app.py
"""

import os
import re
import sys
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import date

try:
    from tkcalendar import DateEntry
    HAS_TKCALENDAR = True
except ImportError:
    HAS_TKCALENDAR = False

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import autocomplete_db


# ---------------------------------------------------------------------------
# Field configuration
# ---------------------------------------------------------------------------
# Each entry: (key, label, widget_type, options)
# widget_type is one of: "text", "date", "dropdown", "paragraph"

TEST_OPTIONS = [
    "PSG",
    "SPLIT PSG",
    "Titration",
    "CPAP Titration",
    "BIPAP Titration",
    "BIPAP ST titration",
    "ASV Titration",
    "Pediatric PSG",
]

LAB_LOCATION_OPTIONS = [
    "Loma Linda",
    "Long Beach",
    "Torrance",
    "Palm Desert",
]

SLEEP_COMPLAINTS_OPTIONS = [
    "Dry Mouth",
    "Nasal Congestion",
    "Headache",
    "Heartburn",
    "Chest Pain",
    "Choking and Gasping",
]

ROOM_OPTIONS = [str(n) for n in range(1, 7)]  # "1" through "6"

FIELDS = [
    ("study_date", "Date of Study", "date", None),
    ("lab_location", "Lab Location", "dropdown", LAB_LOCATION_OPTIONS),
    ("patient_name", "Patient Name", "text", None),
    ("dob", "Date of Birth", "date", None),
    ("room_number", "Room #", "dropdown", ROOM_OPTIONS),
    ("patient_id", "Patient ID", "text", None),
    ("acquisition_num", "Acquisition #", "text", None),
    ("test_ordered", "Test Ordered", "dropdown", TEST_OPTIONS),
    ("test_conducted", "Test Conducted", "dropdown", TEST_OPTIONS),
    ("mask_used", "Mask Used", "text", None),
    ("ordering_md", "Ordering MD", "text", None),
    ("recording_tech", "Recording Tech", "text", None),
    ("height", "Height (inches)", "text", None),
    ("weight", "Weight (lbs.)", "text", None),
    ("bmi", "BMI (Kg/m\u2082)", "text", None),
    ("epworth", "Epworth", "text", None),
    ("lights_off", "Lights Off", "text", None),
    ("lights_on", "Lights On", "text", None),
    ("sleep_complaints", "Sleep Complaints", "multiselect", SLEEP_COMPLAINTS_OPTIONS),
    ("medical_history", "Medical History", "text", None),
    ("medication_list", "Medication List", "text", None),
]

REQUIRED_FIELDS = {"patient_name", "patient_id"}

# Fields backed by the SQLite autocomplete database (see
# autocomplete_db.py). Suggestions grow over time as notes are saved,
# rather than being a fixed hard-coded list.
#
# multi_value=True fields are comma-separated lists (autocomplete
# matches whatever's typed after the last comma). multi_value=False
# fields hold a single value that may itself legitimately contain a
# comma (e.g. "Smith, John MD"), so the whole field is treated as one
# token instead of being split.
DB_AUTOCOMPLETE_FIELDS = {
    "mask_used": {"multi_value": False},
    "ordering_md": {"multi_value": False},
    "recording_tech": {"multi_value": False},
    "medical_history": {"multi_value": True},
    "medication_list": {"multi_value": True},
}

# Starter suggestions seeded into the database the very first time the
# app runs (only for a field that has no data yet -- never overwrites
# real entered data on later runs).
SEED_SUGGESTIONS = {
    "medical_history": [
        "Hypertension", "Type 2 Diabetes", "Type 1 Diabetes", "Obesity",
        "GERD", "Atrial Fibrillation", "COPD", "Asthma",
        "Coronary Artery Disease", "Congestive Heart Failure",
        "Depression", "Anxiety", "Hypothyroidism", "Hyperthyroidism",
        "Stroke / CVA", "Chronic Kidney Disease", "Prior Sleep Apnea Diagnosis",
        "Insomnia", "Restless Leg Syndrome", "Narcolepsy", "Fibromyalgia",
        "Chronic Pain", "Seizure Disorder", "Migraine", "Osteoarthritis",
        "Hyperlipidemia", "Anemia", "PTSD",
    ],
    "medication_list": [
        "Lisinopril", "Metformin", "Atorvastatin", "Levothyroxine",
        "Metoprolol", "Losartan", "Albuterol", "Omeprazole", "Gabapentin",
        "Trazodone", "Zolpidem (Ambien)", "Melatonin", "Sertraline",
        "Escitalopram", "Fluoxetine", "Clonazepam", "Alprazolam",
        "Amlodipine", "Hydrochlorothiazide", "Furosemide", "Insulin",
        "Warfarin", "Apixaban", "Prednisone", "Montelukast", "Duloxetine",
        "Bupropion", "Venlafaxine", "Amitriptyline", "Modafinil",
        "Armodafinil", "Pantoprazole", "Rosuvastatin", "Simvastatin",
    ],
}


# ---------------------------------------------------------------------------
# Scrollable frame helper
# ---------------------------------------------------------------------------
class ScrollableFrame(ttk.Frame):
    """A frame with a vertical scrollbar; put widgets in .inner"""

    def __init__(self, parent, *args, **kwargs):
        super().__init__(parent, *args, **kwargs)

        canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0)
        vscroll = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.inner = ttk.Frame(canvas)

        # The window this canvas embeds. We keep the id so we can force
        # its width to track the canvas's width below -- without this,
        # some Tk builds (notably on Windows) leave the embedded frame
        # at its default 1x1 size and the canvas renders empty.
        self._window_id = canvas.create_window(
            (0, 0), window=self.inner, anchor="nw"
        )

        def _on_inner_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def _on_canvas_configure(event):
            # Force the embedded frame to always match the canvas width
            canvas.itemconfig(self._window_id, width=event.width)

        self.inner.bind("<Configure>", _on_inner_configure)
        canvas.bind("<Configure>", _on_canvas_configure)
        canvas.configure(yscrollcommand=vscroll.set)

        canvas.pack(side="left", fill="both", expand=True)
        vscroll.pack(side="right", fill="y")

        # Mouse wheel support (only while the cursor is over this widget,
        # not bound globally, so it doesn't interfere with other windows)
        def _on_mousewheel(event):
            if sys.platform == "darwin":
                # macOS (including trackpad) reports small deltas
                # (typically ±1-3 per scroll tick) -- using them
                # directly feels natural, unlike Windows below.
                delta = -1 * event.delta
            else:
                # Windows reports deltas in multiples of 120 per tick.
                delta = int(-1 * (event.delta / 120))
            canvas.yview_scroll(delta, "units")

        def _on_linux_scroll_up(event):
            canvas.yview_scroll(-1, "units")

        def _on_linux_scroll_down(event):
            canvas.yview_scroll(1, "units")

        def _bind_wheel(event):
            canvas.bind_all("<MouseWheel>", _on_mousewheel)
            canvas.bind_all("<Button-4>", _on_linux_scroll_up)
            canvas.bind_all("<Button-5>", _on_linux_scroll_down)

        def _unbind_wheel(event):
            canvas.unbind_all("<MouseWheel>")
            canvas.unbind_all("<Button-4>")
            canvas.unbind_all("<Button-5>")

        canvas.bind("<Enter>", _bind_wheel)
        canvas.bind("<Leave>", _unbind_wheel)

        # Force an initial layout pass so content is visible immediately
        # instead of waiting on the first natural <Configure> event.
        self.after(0, lambda: canvas.event_generate("<Configure>"))


# ---------------------------------------------------------------------------
# Autocomplete entry widget
# ---------------------------------------------------------------------------
class AutocompleteEntry(ttk.Entry):
    """A ttk.Entry that shows a dropdown of matching suggestions as you
    type.

    With multi_value=True (the default), it's designed for comma-
    separated free-text fields: it matches and completes whatever
    you're currently typing after the last comma, leaving earlier
    entries in the field untouched, and appends ", " after a selection
    so you're ready to type the next item.

    With multi_value=False, the whole field is treated as a single
    value (matching is done against the entire field content, not
    split on commas -- important for fields like a doctor's name that
    may legitimately contain a comma, e.g. "Smith, John MD"), and
    selecting a suggestion replaces the whole field rather than
    appending.

    Click a suggestion, or use Down/Up + Enter, to accept it. Escape
    closes the dropdown without selecting anything.

    The dropdown is a Listbox placed (via .place()) directly on top of
    the entry's own top-level window, rather than a separate Toplevel
    popup. A second real OS-level window is what caused a focus/grab
    freeze on macOS -- keeping everything inside one window avoids that
    entirely.
    """

    MAX_SUGGESTIONS = 8
    ROW_HEIGHT = 20

    def __init__(self, master, suggestions=None, multi_value=True, **kwargs):
        super().__init__(master, **kwargs)
        self.suggestions = suggestions or []
        self.multi_value = multi_value
        self._listbox = None  # created lazily, reused (shown/hidden)

        self.bind("<KeyRelease>", self._on_key_release)
        self.bind("<Down>", self._focus_listbox)
        self.bind("<Escape>", lambda e: self._hide_popup())
        self.bind("<FocusOut>", self._on_focus_out)

    # -- suggestion matching -----------------------------------------
    def _current_token(self):
        """Text currently being typed. For multi_value fields, that's
        whatever's after the last comma; for single-value fields, it's
        the entire field content."""
        text = self.get()
        if not self.multi_value:
            return text.strip(), -1
        last_comma = text.rfind(",")
        token = text[last_comma + 1:] if last_comma != -1 else text
        return token.strip(), last_comma

    def _matches(self, token):
        token_lower = token.lower()
        starts = [s for s in self.suggestions if s.lower().startswith(token_lower)]
        contains = [
            s for s in self.suggestions
            if token_lower in s.lower() and s not in starts
        ]
        return (starts + contains)[: self.MAX_SUGGESTIONS]

    # -- event handlers -------------------------------------------------
    def _on_key_release(self, event):
        if event.keysym in (
            "Down", "Up", "Return", "Escape", "Tab",
            "Shift_L", "Shift_R", "Control_L", "Control_R",
        ):
            return
        token, _ = self._current_token()
        if not token:
            self._hide_popup()
            return
        matches = self._matches(token)
        if matches:
            self._show_popup(matches)
        else:
            self._hide_popup()

    def _on_focus_out(self, event):
        # Delay so a click on the listbox registers before we hide it.
        self.after(150, self._hide_popup_if_unfocused)

    def _hide_popup_if_unfocused(self):
        focused = self.focus_get()
        if focused is not self and focused is not self._listbox:
            self._hide_popup()

    def _focus_listbox(self, event):
        if self._listbox is not None and self._listbox.winfo_ismapped():
            self._listbox.focus_set()
            self._listbox.selection_set(0)
            self._listbox.activate(0)
        return "break"

    # -- popup management -------------------------------------------
    def _ensure_listbox(self):
        if self._listbox is None:
            toplevel = self.winfo_toplevel()
            self._listbox = tk.Listbox(
                toplevel, activestyle="dotbox", exportselection=False,
                relief="solid", borderwidth=1,
            )
            self._listbox.bind("<ButtonRelease-1>", self._select_from_click)
            self._listbox.bind("<Return>", self._select_from_key)
            self._listbox.bind("<Escape>", lambda e: self._hide_popup())
            self._listbox.bind("<FocusOut>", self._on_focus_out)
        return self._listbox

    def _show_popup(self, matches):
        listbox = self._ensure_listbox()
        toplevel = self.winfo_toplevel()

        x = self.winfo_rootx() - toplevel.winfo_rootx()
        y = self.winfo_rooty() - toplevel.winfo_rooty() + self.winfo_height()
        width = max(self.winfo_width(), 180)
        height_px = min(len(matches), self.MAX_SUGGESTIONS) * self.ROW_HEIGHT

        listbox.delete(0, tk.END)
        for m in matches:
            listbox.insert(tk.END, m)

        listbox.place(x=x, y=y, width=width, height=height_px)
        listbox.lift()

    def _hide_popup(self):
        if self._listbox is not None:
            self._listbox.place_forget()

    def _select_from_click(self, event):
        selection = self._listbox.curselection()
        if selection:
            self._apply_selection(self._listbox.get(selection[0]))

    def _select_from_key(self, event):
        selection = self._listbox.curselection()
        if selection:
            self._apply_selection(self._listbox.get(selection[0]))
        return "break"

    def _apply_selection(self, chosen):
        if not self.multi_value:
            new_text = chosen
        else:
            text = self.get()
            _, last_comma = self._current_token()
            prefix = text[: last_comma + 1] + " " if last_comma != -1 else ""
            new_text = f"{prefix}{chosen}, "
        self.delete(0, tk.END)
        self.insert(0, new_text)
        self.icursor(tk.END)
        self._hide_popup()
        self.focus_set()


# ---------------------------------------------------------------------------
# Multi-select combo box (checklist dropdown)
# ---------------------------------------------------------------------------
class MultiSelectCombobox(ttk.Frame):
    """A combobox-like widget for picking multiple values from a fixed
    list (e.g. Sleep Complaints). Selected values show comma-separated
    in a read-only display field; clicking it (or the dropdown button)
    opens a checklist of all options.

    Like AutocompleteEntry's dropdown, the checklist is overlaid on the
    same window via .place() rather than a separate Toplevel, to avoid
    the macOS focus/grab freeze a second real window caused earlier.
    """

    def __init__(self, master, options, **kwargs):
        super().__init__(master, **kwargs)
        self.options = options
        self._vars = {opt: tk.BooleanVar(value=False) for opt in options}
        self._popup = None
        self._click_away_id = None

        self.display = ttk.Entry(self)
        self.display.configure(state="readonly")
        self.display.pack(side="left", fill="both", expand=True)
        self.display.bind("<Button-1>", self._toggle_popup)

        self.toggle_btn = ttk.Button(
            self, text="\u25bc", width=2, command=self._toggle_popup
        )
        self.toggle_btn.pack(side="left")

    # -- public interface (mirrors the bits of Entry the rest of the
    # app relies on: .get() and a way to clear/restore the value) ----
    def get(self):
        return ", ".join(opt for opt in self.options if self._vars[opt].get())

    def clear(self):
        for var in self._vars.values():
            var.set(False)
        self._refresh_display()

    def set_from_text(self, text):
        """Check whichever options appear (comma-separated) in text."""
        chosen = {t.strip() for t in text.split(",") if t.strip()}
        for opt in self.options:
            self._vars[opt].set(opt in chosen)
        self._refresh_display()

    # -- internal ------------------------------------------------------
    def _refresh_display(self):
        self.display.configure(state="normal")
        self.display.delete(0, tk.END)
        self.display.insert(0, self.get())
        self.display.configure(state="readonly")

    def _toggle_popup(self, event=None):
        if self._popup is not None:
            self._hide_popup()
        else:
            self._show_popup()
        return "break"

    def _show_popup(self):
        toplevel = self.winfo_toplevel()
        self._popup = tk.Frame(toplevel, relief="solid", borderwidth=1, bg="white")

        for opt in self.options:
            cb = tk.Checkbutton(
                self._popup, text=opt, variable=self._vars[opt],
                anchor="w", bg="white", activebackground="#e8f0fe",
                highlightthickness=0, command=self._refresh_display,
            )
            cb.pack(fill="x", anchor="w", padx=4)

        x = self.winfo_rootx() - toplevel.winfo_rootx()
        y = self.winfo_rooty() - toplevel.winfo_rooty() + self.winfo_height()
        width = max(self.winfo_width(), 220)
        self._popup.place(x=x, y=y, width=width)
        self._popup.lift()

        # Any click outside the popup (and outside the display/button
        # that opened it) should close it.
        self._click_away_id = toplevel.bind(
            "<Button-1>", self._on_click_away, add="+"
        )

    def _on_click_away(self, event):
        if self._popup is None:
            return
        widget = event.widget
        if widget in (self.toggle_btn, self.display):
            return
        node = widget
        while node is not None:
            if node == self._popup:
                return
            node = node.master
        self._hide_popup()

    def _hide_popup(self):
        if self._popup is not None:
            toplevel = self.winfo_toplevel()
            if self._click_away_id is not None:
                toplevel.unbind("<Button-1>", self._click_away_id)
                self._click_away_id = None
            self._popup.destroy()
            self._popup = None


# ---------------------------------------------------------------------------
# Tech Note entry window
# ---------------------------------------------------------------------------
class TechNoteWindow(tk.Toplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("New Tech Note")
        self.geometry("620x720")
        self.minsize(520, 400)

        self.widgets = {}  # key -> widget
        self.user_narratives_menu = None  # created lazily, once a narrative exists
        self._build_menu_bar()

        container = ttk.Frame(self, padding=10)
        container.pack(fill="both", expand=True)

        header = ttk.Label(
            container, text="Sleep Study Tech Note", font=("Segoe UI", 16, "bold")
        )
        header.pack(anchor="w", pady=(0, 10))

        scroll_area = ScrollableFrame(container)
        scroll_area.pack(fill="both", expand=True)
        form = scroll_area.inner

        form.columnconfigure(1, weight=1)

        row = 0
        for key, label, wtype, options in FIELDS:
            lbl_text = label + (" *" if key in REQUIRED_FIELDS else "")
            ttk.Label(form, text=lbl_text).grid(
                row=row, column=0, sticky="w", padx=(0, 10), pady=4
            )

            if key == "patient_id":
                # Auto-generated (Date of Study + Room #), never typed
                # directly -- see _update_patient_id().
                widget = ttk.Entry(form)
                widget.configure(state="readonly")
            elif wtype == "date":
                if HAS_TKCALENDAR:
                    widget = DateEntry(
                        form, date_pattern="mm/dd/yyyy", width=18,
                        maxdate=date.today(),
                    )
                else:
                    widget = ttk.Entry(form, width=20)
                    widget.insert(0, "MM/DD/YYYY")
                if key == "study_date":
                    widget.bind("<<DateEntrySelected>>", self._update_patient_id)
                    widget.bind("<KeyRelease>", self._update_patient_id)
                    widget.bind("<FocusOut>", self._update_patient_id)
            elif wtype == "dropdown":
                widget = ttk.Combobox(
                    form, values=options, state="readonly", width=25
                )
                if key == "room_number":
                    widget.bind("<<ComboboxSelected>>", self._update_patient_id)
            elif wtype == "multiselect":
                widget = MultiSelectCombobox(form, options=options)
            elif key in DB_AUTOCOMPLETE_FIELDS:
                config = DB_AUTOCOMPLETE_FIELDS[key]
                widget = AutocompleteEntry(
                    form,
                    suggestions=autocomplete_db.get_suggestions(key),
                    multi_value=config["multi_value"],
                )
            else:  # plain text
                widget = ttk.Entry(form)

            widget.grid(row=row, column=1, sticky="ew", pady=4)
            self.widgets[key] = widget

            # Small helper button next to BMI to auto-calc from height/weight
            if key == "bmi":
                calc_btn = ttk.Button(
                    form, text="Calc", width=6,
                    command=self._calc_bmi,
                )
                calc_btn.grid(row=row, column=2, padx=(6, 0))

            row += 1

        # Now that study_date and room_number widgets both exist,
        # populate Patient ID from their current values (study_date
        # defaults to today, so this only needs a room selected).
        self._update_patient_id()

        # Tech comment - big paragraph box
        comment_header = ttk.Frame(form)
        comment_header.grid(row=row, column=0, columnspan=3, sticky="ew", pady=(14, 4))
        comment_header.columnconfigure(0, weight=1)

        ttk.Label(
            comment_header, text="Tech Comment", font=("Segoe UI", 10, "bold")
        ).grid(row=0, column=0, sticky="w")
        ttk.Button(
            comment_header, text="Clear", width=8, command=self._clear_tech_comment
        ).grid(row=0, column=1, sticky="e")
        row += 1

        comment_box = tk.Text(form, height=10, wrap="word", undo=True)
        comment_box.grid(row=row, column=0, columnspan=3, sticky="nsew", pady=(0, 6))
        form.rowconfigure(row, weight=1)
        self.widgets["tech_comment"] = comment_box
        row += 1

        note = ttk.Label(
            container, text="* Required fields", foreground="gray"
        )
        note.pack(anchor="w", pady=(4, 6))

        # Buttons
        btn_frame = ttk.Frame(container)
        btn_frame.pack(fill="x", pady=(4, 0))

        ttk.Button(btn_frame, text="Clear", command=self._clear_form).pack(
            side="left"
        )
        ttk.Button(btn_frame, text="Cancel", command=self.destroy).pack(
            side="right", padx=(6, 0)
        )
        ttk.Button(
            btn_frame, text="Save Tech Note", command=self._save
        ).pack(side="right")

    # -- tech narratives -----------------------------------------------
    def _build_menu_bar(self):
        menu_bar = tk.Menu(self, tearoff=0)
        self.config(menu=menu_bar)
        self.menu_bar = menu_bar
        menu_bar.add_command(label="Add Tech Narrative", command=self._add_tech_narrative)
        self._refresh_user_narratives_menu()

    def _refresh_user_narratives_menu(self):
        """(Re)build the 'User Narratives' dropdown from the database.
        The dropdown only appears in the menu bar once at least one
        narrative has been saved."""
        narratives = autocomplete_db.get_narratives()
        if not narratives:
            return

        if self.user_narratives_menu is None:
            self.user_narratives_menu = tk.Menu(self.menu_bar, tearoff=0)
            self.menu_bar.add_cascade(label="User Narratives", menu=self.user_narratives_menu)
        else:
            self.user_narratives_menu.delete(0, tk.END)

        for name, text in narratives:
            self.user_narratives_menu.add_command(
                label=name, command=lambda t=text: self._insert_narrative(t)
            )

    def _insert_narrative(self, text):
        comment_box = self.widgets.get("tech_comment")
        if comment_box is None:
            return
        comment_box.insert(tk.INSERT, text)
        comment_box.focus_set()

    def _clear_tech_comment(self):
        comment_box = self.widgets.get("tech_comment")
        if comment_box is None:
            return
        comment_box.delete("1.0", tk.END)
        comment_box.focus_set()

    def _add_tech_narrative(self):
        dialog = tk.Toplevel(self)
        dialog.title("Add Tech Narrative")
        dialog.geometry("440x380")
        dialog.transient(self)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding=10)
        frame.pack(fill="both", expand=True)

        ttk.Label(frame, text="Narrative Name").pack(anchor="w")
        name_entry = ttk.Entry(frame)
        name_entry.pack(fill="x", pady=(0, 10))

        ttk.Label(frame, text="Narrative Text").pack(anchor="w")
        text_box = tk.Text(frame, height=12, wrap="word", undo=True)
        text_box.pack(fill="both", expand=True, pady=(0, 10))

        def on_save():
            name = name_entry.get().strip()
            text = text_box.get("1.0", tk.END).strip()
            if not name or not text:
                messagebox.showerror(
                    "Missing Information",
                    "Please enter both a name and the narrative text.",
                    parent=dialog,
                )
                return
            autocomplete_db.save_narrative(name, text)
            self._refresh_user_narratives_menu()
            dialog.destroy()

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill="x")
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(
            side="right", padx=(6, 0)
        )
        ttk.Button(btn_frame, text="Save Narrative", command=on_save).pack(side="right")

        name_entry.focus_set()

    # -- helpers ------------------------------------------------------
    def _update_patient_id(self, event=None):
        """Recompute Patient ID from Date of Study + Room #, formatted
        as MMDDYYYY-<room>. Leaves it blank until both are set to a
        complete, valid value."""
        date_widget = self.widgets.get("study_date")
        room_widget = self.widgets.get("room_number")
        if date_widget is None or room_widget is None:
            return

        date_text = date_widget.get().strip()
        room_text = room_widget.get().strip()

        if re.match(r"^\d{2}/\d{2}/\d{4}$", date_text) and room_text:
            new_id = f"{date_text.replace('/', '')}-RM{room_text}"
        else:
            new_id = ""

        self._set_patient_id_text(new_id)

    def _set_patient_id_text(self, text):
        """Patient ID is a read-only Entry -- readonly Entries reject
        programmatic edits too, so briefly toggle to normal to update
        it, then lock it back down."""
        widget = self.widgets.get("patient_id")
        if widget is None:
            return
        widget.configure(state="normal")
        widget.delete(0, tk.END)
        widget.insert(0, text)
        widget.configure(state="readonly")

    def _calc_bmi(self):
        """Auto-calculate BMI from Height (in inches) and Weight (in lb)."""
        try:
            height_in = float(self.widgets["height"].get().strip())
            weight_lb = float(self.widgets["weight"].get().strip())
            bmi = (weight_lb / (height_in ** 2)) * 703
            bmi_widget = self.widgets["bmi"]
            bmi_widget.delete(0, tk.END)
            bmi_widget.insert(0, f"{bmi:.1f}")
        except (ValueError, ZeroDivisionError):
            messagebox.showwarning(
                "Cannot Calculate BMI",
                "Enter numeric Height (inches) and Weight (lb) first.",
            )

    def _clear_form(self):
        if not messagebox.askyesno("Clear Form", "Clear all fields?"):
            return
        for key, widget in self.widgets.items():
            if key == "patient_id":
                continue  # handled below, after study_date/room_number reset
            elif isinstance(widget, MultiSelectCombobox):
                widget.clear()
            elif isinstance(widget, ttk.Combobox):
                widget.set("")
            elif isinstance(widget, tk.Text):
                widget.delete("1.0", tk.END)
            elif HAS_TKCALENDAR and isinstance(widget, DateEntry):
                widget.set_date(date.today())
            else:
                widget.delete(0, tk.END)
        # Room # is now cleared and study_date reset to today -- recompute
        # Patient ID (will end up blank, since Room # is required for it).
        self._update_patient_id()

    def _collect_values(self):
        values = {}
        for key, widget in self.widgets.items():
            if isinstance(widget, tk.Text):
                values[key] = widget.get("1.0", tk.END).strip()
            elif HAS_TKCALENDAR and isinstance(widget, DateEntry):
                values[key] = widget.get()
            else:
                values[key] = widget.get().strip()
        return values

    def _save(self):
        values = self._collect_values()

        missing = [
            label for key, label, *_ in FIELDS
            if key in REQUIRED_FIELDS and not values.get(key)
        ]
        if missing:
            messagebox.showerror(
                "Missing Information",
                "Please fill in: " + ", ".join(missing),
            )
            return

        default_name = self._suggest_filename(values)
        path = filedialog.asksaveasfilename(
            title="Save Tech Note",
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word Document", "*.docx")],
        )
        if not path:
            return

        try:
            build_docx(values, path)
        except Exception as exc:  # noqa: BLE001
            messagebox.showerror("Error Saving Document", str(exc))
            return

        self._record_autocomplete_values(values)

        messagebox.showinfo("Saved", f"Tech note saved to:\n{path}")
        self.destroy()

    @staticmethod
    def _record_autocomplete_values(values):
        """Feed whatever was entered in the DB-backed fields back into
        the database, so future notes get better suggestions."""
        for key, config in DB_AUTOCOMPLETE_FIELDS.items():
            text = values.get(key, "")
            if not text:
                continue
            if config["multi_value"]:
                items = [t.strip() for t in text.split(",") if t.strip()]
                autocomplete_db.record_values(key, items)
            else:
                autocomplete_db.record_value(key, text.strip())

    @staticmethod
    def _suggest_filename(values):
        name = values.get("patient_name", "").strip() or "TechNote"
        safe = re.sub(r"[^A-Za-z0-9_\- ]", "", name).strip().replace(" ", "_")
        today_str = date.today().strftime("%Y-%m-%d")
        return f"{safe}_TechNote_{today_str}.docx"


# ---------------------------------------------------------------------------
# Word document generation
# ---------------------------------------------------------------------------
def build_docx(values, path):
    values = dict(values)  # work on a copy -- don't mutate the caller's data
    epworth_raw = values.get("epworth", "").strip()
    if epworth_raw and "/" not in epworth_raw:
        values["epworth"] = f"{epworth_raw}/24"

    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Sleep Study Tech Note", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated = doc.add_paragraph(
        f"Generated: {date.today().strftime('%m/%d/%Y')}"
    )
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.runs[0].italic = True
    generated.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # --- Demographics / study info table -----------------------------
    demo_keys = [
        "study_date", "lab_location",
        "patient_name", "dob", "room_number", "patient_id", "acquisition_num",
        "test_ordered", "test_conducted", "mask_used",
        "ordering_md", "recording_tech",
    ]
    _add_section_heading(doc, "Patient & Study Information")
    _add_field_table(doc, demo_keys, values=values)

    # --- Clinical measures table ---------------------------------------
    clinical_keys = [
        "height", "weight", "bmi", "epworth", "lights_off", "lights_on",
    ]
    _add_section_heading(doc, "Clinical Measures")
    _add_field_table(doc, clinical_keys, values=values)

    # --- Narrative fields ------------------------------------------------
    narrative_keys = [
        ("sleep_complaints", "Sleep Complaints"),
        ("medical_history", "Medical History"),
        ("medication_list", "Medication List"),
    ]
    _add_section_heading(doc, "Clinical History")
    for key, label in narrative_keys:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(values.get(key, "") or "N/A")

    # --- Tech comment -----------------------------------------------
    _add_section_heading(doc, "Tech Comment")
    doc.add_paragraph(values.get("tech_comment", "") or "N/A")

    doc.save(path)
    return path


def _add_section_heading(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(13)


def _label_for(key):
    for k, label, *_ in FIELDS:
        if k == key:
            return label
    return key


def _add_field_table(doc, keys, values=None):
    """Adds a 2-column table of label/value pairs for the given keys."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.2)

    for key in keys:
        row_cells = table.add_row().cells
        row_cells[0].text = _label_for(key)
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[1].text = (values or {}).get(key, "") or ""

    doc.add_paragraph()  # spacer after table


# ---------------------------------------------------------------------------
# Home window
# ---------------------------------------------------------------------------
class HomeWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Sleep Lab Tech Notes")
        self.geometry("420x300")
        self.resizable(False, False)

        if not HAS_TKCALENDAR:
            self.after(200, lambda: messagebox.showwarning(
                "Optional Dependency Missing",
                "tkcalendar is not installed, so the Date of Birth field\n"
                "will be a plain text box instead of a date picker.\n\n"
                "Install it with:  pip install tkcalendar",
            ))

        frame = ttk.Frame(self, padding=30)
        frame.pack(fill="both", expand=True)

        ttk.Label(
            frame, text="Sleep Lab Tech Notes",
            font=("Segoe UI", 18, "bold")
        ).pack(pady=(0, 6))
        ttk.Label(
            frame, text="Create and export tech notes to Word (.docx)",
            foreground="gray"
        ).pack(pady=(0, 30))

        big_button = tk.Button(
            frame,
            text="+ New Tech Note",
            font=("Segoe UI", 13, "bold"),
            bg="#2c6fbb",
            fg="white",
            activebackground="#25599a",
            activeforeground="white",
            relief="flat",
            padx=20,
            pady=14,
            command=self.open_new_tech_note,
        )
        big_button.pack(fill="x")

        ttk.Button(frame, text="Quit", command=self.destroy).pack(pady=(20, 0))

    def open_new_tech_note(self):
        try:
            TechNoteWindow(self)
        except Exception as exc:  # noqa: BLE001
            import traceback
            details = traceback.format_exc()
            print(details)  # also visible in the terminal
            messagebox.showerror(
                "Error Opening Tech Note",
                f"Something went wrong building the form:\n\n{exc}\n\n"
                "See the terminal for full details.",
            )


if __name__ == "__main__":
    autocomplete_db.init_db(seed_defaults=SEED_SUGGESTIONS)
    app = HomeWindow()
    app.mainloop()